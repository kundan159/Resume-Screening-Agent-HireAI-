#!/usr/bin/env python3
"""
Resume Screening Agent
=======================
Parses a folder of resumes (PDF / DOCX / TXT), scores each one against a
Job Description (JD), and outputs a ranked, reasoned shortlist as CSV + JSON.

Usage:
    python3 agent.py --jd job_description.txt --resumes resumes/ --out output/

See SCORING_METHOD.md for a full explanation of how the score is computed.
"""

import argparse
import csv
import json
import os
import re
import sys
from pathlib import Path

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx  # python-docx
except ImportError:
    docx = None


# --------------------------------------------------------------------------
# 1. FILE READING (PDF / DOCX / TXT)
# --------------------------------------------------------------------------

def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_pdf(path: Path) -> str:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is required to read PDF files.")
    text_chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_chunks.append(page.extract_text() or "")
    return "\n".join(text_chunks)


def read_docx(path: Path) -> str:
    if docx is None:
        raise RuntimeError("python-docx is required to read DOCX files.")
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def read_resume(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return read_txt(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix in (".docx",):
        return read_docx(path)
    raise ValueError(f"Unsupported file type: {suffix} ({path.name})")


# --------------------------------------------------------------------------
# 2. JOB DESCRIPTION PARSING
# --------------------------------------------------------------------------

SECTION_HEADERS = ["required skills", "preferred skills", "education", "responsibilities"]

# Canonical taxonomy of atomic skill/technology terms the agent knows how to
# recognize inside free-text JD bullets and resume text. This keeps extracted
# skills as clean single concepts (e.g. "PostgreSQL") instead of whole
# sentences. Extend this list to adapt the agent to other job families.
SKILL_TAXONOMY = [
    "Python", "Java", "Go", "JavaScript", "TypeScript", "Node.js",
    "Flask", "FastAPI", "Django", "Spring Boot", "Express",
    "REST API", "REST APIs", "GraphQL",
    "SQL", "PostgreSQL", "MySQL", "MongoDB",
    "AWS", "EC2", "S3", "Lambda", "RDS", "Azure", "GCP",
    "Docker", "Kubernetes", "Terraform", "CI/CD", "Jenkins", "Git",
    "Kafka", "RabbitMQ", "SQS", "Redis",
    "pytest", "unit testing", "integration testing", "unit tests",
    "system design", "distributed systems",
    "agile", "scrum", "mentoring",
]


def find_taxonomy_terms(text: str) -> list:
    """Return taxonomy terms that appear in the given text, in taxonomy order."""
    text_lower = text.lower()
    found = []
    for term in SKILL_TAXONOMY:
        term_lower = term.lower()
        if re.fullmatch(r"[a-z0-9 ]+", term_lower):
            pattern = r"\b" + re.escape(term_lower) + r"\b"
            hit = re.search(pattern, text_lower)
        else:
            hit = term_lower in text_lower
        if hit and term not in found:
            found.append(term)
    return found


def parse_jd(jd_text: str) -> dict:
    lines = jd_text.splitlines()
    required_skills, preferred_skills = [], []
    current_section = None

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        lower = line.lower().rstrip(":")
        if lower in SECTION_HEADERS:
            current_section = lower
            continue
        if line.startswith("-"):
            content = line.lstrip("- ").strip()
            if current_section == "required skills":
                required_skills.extend(find_taxonomy_terms(content))
            elif current_section == "preferred skills":
                preferred_skills.extend(find_taxonomy_terms(content))

    # Required experience range, e.g. "Experience Required: 4-8 years"
    exp_match = re.search(r"experience required:\s*(\d+)\s*-\s*(\d+)\s*years?", jd_text, re.I)
    min_exp, max_exp = (int(exp_match.group(1)), int(exp_match.group(2))) if exp_match else (0, None)

    return {
        "raw_text": jd_text,
        "required_skills": sorted(set(required_skills), key=str.lower),
        "preferred_skills": sorted(set(preferred_skills), key=str.lower),
        "min_experience": min_exp,
        "max_experience": max_exp,
    }


# --------------------------------------------------------------------------
# 3. RESUME FIELD EXTRACTION
# --------------------------------------------------------------------------

DEGREE_KEYWORDS = [
    "phd", "ph.d", "m.tech", "mtech", "m.s.", "m.sc", "msc", "mba",
    "b.tech", "btech", "b.e.", "be ", "b.sc", "bsc", "bachelor", "master",
]


def guess_name(text: str, fallback: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if line:
            # Heuristic: first non-empty line, if short and has no digits/@ symbol
            if len(line) <= 60 and "@" not in line and not re.search(r"\d", line):
                return line
            break
    return fallback


def extract_experience_years(text: str) -> float:
    # Preferred: explicit "Total years of professional experience: X years"
    m = re.search(r"total years of professional experience:\s*([\d.]+)\s*years?", text, re.I)
    if m:
        return float(m.group(1))

    # Fallback: sum bracketed durations like "[3 years]" or "[6 months]"
    total = 0.0
    for val, unit in re.findall(r"\[([\d.]+)\s*(year|years|month|months)\]", text, re.I):
        val = float(val)
        total += val if "year" in unit.lower() else val / 12.0
    if total > 0:
        return round(total, 1)

    # Last resort: "X years of experience" anywhere in the text
    m = re.search(r"([\d.]+)\+?\s*years? of experience", text, re.I)
    if m:
        return float(m.group(1))

    return 0.0


def extract_education(text: str) -> str:
    lines = text.splitlines()
    edu_section = False
    matches = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.lower().rstrip(":") == "education":
            edu_section = True
            continue
        if edu_section:
            if stripped.lower().rstrip(":") in SECTION_HEADERS or stripped.isupper():
                edu_section = False
                continue
            if any(kw in stripped.lower() for kw in DEGREE_KEYWORDS):
                matches.append(stripped)
    if matches:
        return matches[0]

    # Fallback: search whole document for a degree-keyword line
    for line in lines:
        if any(kw in line.lower() for kw in DEGREE_KEYWORDS):
            return line.strip()
    return "Not found"


def extract_matched_skills(text: str, skill_list):
    text_lower = text.lower()
    matched = []
    for skill in skill_list:
        skill_lower = skill.lower()
        if re.fullmatch(r"[a-z0-9 ]+", skill_lower):
            pattern = r"\b" + re.escape(skill_lower) + r"\b"
            hit = re.search(pattern, text_lower)
        else:
            hit = skill_lower in text_lower
        if hit:
            matched.append(skill)
    return matched


# --------------------------------------------------------------------------
# 4. SCORING
# --------------------------------------------------------------------------

def experience_score(candidate_years: float, min_exp: int, max_exp) -> float:
    """1.0 if candidate meets/exceeds the required range, partial credit below it."""
    if min_exp == 0:
        return 1.0
    if candidate_years >= min_exp:
        return 1.0
    # Linear partial credit for under-experienced candidates
    return max(0.0, candidate_years / min_exp)


def score_candidates(jd: dict, resumes: list) -> list:
    corpus = [jd["raw_text"]] + [r["raw_text"] for r in resumes]
    vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
    tfidf = vectorizer.fit_transform(corpus)
    jd_vec, resume_vecs = tfidf[0:1], tfidf[1:]
    similarities = cosine_similarity(jd_vec, resume_vecs)[0]

    all_skills = jd["required_skills"] + jd["preferred_skills"]

    results = []
    for resume, sim in zip(resumes, similarities):
        matched_required = extract_matched_skills(resume["raw_text"], jd["required_skills"])
        matched_preferred = extract_matched_skills(resume["raw_text"], jd["preferred_skills"])
        missing_required = [s for s in jd["required_skills"] if s not in matched_required]

        req_ratio = len(matched_required) / len(jd["required_skills"]) if jd["required_skills"] else 1.0
        pref_ratio = len(matched_preferred) / len(jd["preferred_skills"]) if jd["preferred_skills"] else 0.0
        skill_score = 0.8 * req_ratio + 0.2 * pref_ratio

        exp_score = experience_score(resume["experience_years"], jd["min_experience"], jd["max_experience"])
        semantic_score = float(sim)

        final_score = (0.40 * semantic_score) + (0.40 * skill_score) + (0.20 * exp_score)

        reasoning = (
            f"Matched {len(matched_required)}/{len(jd['required_skills'])} required skills "
            f"({', '.join(matched_required) if matched_required else 'none'}); "
            f"{len(matched_preferred)}/{len(jd['preferred_skills'])} preferred skills matched. "
            f"{resume['experience_years']} yrs experience vs. "
            f"{jd['min_experience']}+ yrs required. "
            f"Missing required skills: {', '.join(missing_required) if missing_required else 'none'}."
        )

        results.append({
            "candidate_name": resume["name"],
            "file": resume["file"],
            "final_score": round(final_score * 100, 1),
            "semantic_similarity": round(semantic_score * 100, 1),
            "skill_match_score": round(skill_score * 100, 1),
            "experience_score": round(exp_score * 100, 1),
            "experience_years": resume["experience_years"],
            "education": resume["education"],
            "matched_required_skills": matched_required,
            "missing_required_skills": missing_required,
            "matched_preferred_skills": matched_preferred,
            "reasoning": reasoning,
        })

    results.sort(key=lambda r: r["final_score"], reverse=True)
    for i, r in enumerate(results, start=1):
        r["rank"] = i
    return results


# --------------------------------------------------------------------------
# 5. MAIN PIPELINE
# --------------------------------------------------------------------------

def load_resumes(folder: Path) -> list:
    resumes = []
    supported = (".txt", ".pdf", ".docx")
    for path in sorted(folder.iterdir()):
        if path.suffix.lower() not in supported:
            continue
        try:
            raw_text = read_resume(path)
        except Exception as e:
            print(f"  [warn] Skipping {path.name}: {e}", file=sys.stderr)
            continue
        if not raw_text.strip():
            print(f"  [warn] No extractable text in {path.name}, skipping.", file=sys.stderr)
            continue
        resumes.append({
            "file": path.name,
            "raw_text": raw_text,
            "name": guess_name(raw_text, fallback=path.stem.replace("_", " ").title()),
            "experience_years": extract_experience_years(raw_text),
            "education": extract_education(raw_text),
        })
    return resumes


def write_outputs(results: list, out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)

    json_path = out_dir / "ranked_candidates.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2)

    csv_path = out_dir / "ranked_candidates.csv"
    fieldnames = [
        "rank", "candidate_name", "file", "final_score", "semantic_similarity",
        "skill_match_score", "experience_score", "experience_years", "education",
        "matched_required_skills", "missing_required_skills",
        "matched_preferred_skills", "reasoning",
    ]
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for r in results:
            row = dict(r)
            for key in ("matched_required_skills", "missing_required_skills", "matched_preferred_skills"):
                row[key] = "; ".join(row[key])
            writer.writerow(row)

    return json_path, csv_path


def main():
    parser = argparse.ArgumentParser(description="Rank resumes against a job description.")
    parser.add_argument("--jd", required=True, help="Path to the job description text file.")
    parser.add_argument("--resumes", required=True, help="Folder containing resume files (.pdf/.docx/.txt).")
    parser.add_argument("--out", default="output", help="Output folder for ranked results.")
    args = parser.parse_args()

    jd_text = read_txt(Path(args.jd))
    jd = parse_jd(jd_text)

    print(f"Parsed JD -> {len(jd['required_skills'])} required skills, "
          f"{len(jd['preferred_skills'])} preferred skills, "
          f"{jd['min_experience']}+ years required.")

    resumes_folder = Path(args.resumes)
    resumes = load_resumes(resumes_folder)
    print(f"Loaded {len(resumes)} resumes from {resumes_folder}/")

    results = score_candidates(jd, resumes)

    out_dir = Path(args.out)
    json_path, csv_path = write_outputs(results, out_dir)
    print(f"Wrote {json_path} and {csv_path}")

    print("\nShortlist (top candidates):")
    for r in results:
        print(f"  #{r['rank']:>2}  {r['final_score']:>5.1f}  {r['candidate_name']} ({r['file']})")


if __name__ == "__main__":
    main()
